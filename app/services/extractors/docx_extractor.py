import docx
from .base_extractor import BaseExtractor


class DocxExtractor(BaseExtractor):

    def extract(self, path: str) -> str:
        try:
            document = docx.Document(path)

            parts = []

            # ---- Extract paragraphs ----
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    parts.append(paragraph.text)

            # ---- Extract tables ----
            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        parts.append(" | ".join(row_text))

            # ---- Combine all extracted text ----
            text = "\n".join(parts).strip()

            return text if text else "[DOCX contains no readable text]"

        except Exception as e:
            # Optional: return structured error for debugging
            return f"[Failed to extract DOCX: {str(e)}]"
